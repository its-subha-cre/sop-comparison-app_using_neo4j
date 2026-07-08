import os
import pypdf
import docx

class IngestionAgent:
    """Agent responsible for parsing PDF and DOCX documents and extracting text/tables/metadata."""

    def parse_file(self, filepath: str) -> dict:
        """Determines the file type and extracts metadata, text, and structure."""
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf(filepath)
        elif ext in ['.docx', '.doc']:
            return self._parse_docx(filepath)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _parse_pdf(self, filepath: str) -> dict:
        """Parses PDF utilizing pypdf."""
        reader = pypdf.PdfReader(filepath)
        text = ""
        metadata = {
            "title": reader.metadata.title if reader.metadata and reader.metadata.title else "",
            "author": reader.metadata.author if reader.metadata and reader.metadata.author else "",
            "pages": len(reader.pages),
            "filename": os.path.basename(filepath)
        }
        
        for page in reader.pages:
            text_content = page.extract_text()
            if text_content:
                text += text_content + "\n"
            
        return {
            "metadata": metadata,
            "raw_text": text,
            "tables": []  # Stub for advanced table extractor
        }

    def _parse_docx(self, filepath: str) -> dict:
        """Parses DOCX utilizing python-docx."""
        doc = docx.Document(filepath)
        text = ""
        metadata = {
            "title": os.path.basename(filepath),
            "pages": None,  # Word files don't have page counts in standard core properties easily
            "filename": os.path.basename(filepath)
        }
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Optional table extraction stub
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return {
            "metadata": metadata,
            "raw_text": text,
            "tables": tables
        }
